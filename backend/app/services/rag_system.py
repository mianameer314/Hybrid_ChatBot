"""
RAG (Retrieval-Augmented Generation) System
Supports PDFs, web content, and database queries with vector embeddings
"""
import os
import asyncio
import logging
from typing import List, Dict, Any, Optional, Tuple
import hashlib
from datetime import datetime
import json

# Document processing
import PyPDF2
import docx
from bs4 import BeautifulSoup
import requests
from urllib.parse import urljoin, urlparse
import aiohttp
import aiofiles

# Vector embeddings and search
from sentence_transformers import SentenceTransformer
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
import faiss

# Text processing
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

from app.core.config import settings
from app.core.database import get_db_session
from app.models import Document as DBDocument, DocumentChunk, WebSource, WebChunk, KnowledgeBase
from app.core.cache import cache, make_rag_key
from app.services.llm_providers import llm_manager

logger = logging.getLogger(__name__)

class EmbeddingManager:
    """Manages text embeddings using sentence-transformers"""
    
    def __init__(self, model_name: str = None):
        self.model_name = model_name or settings.EMBEDDING_MODEL
        self.model = None
        self.dimension = settings.VECTOR_DIMENSION
        self.is_initialized = False
    
    async def initialize(self):
        """Initialize the embedding model"""
        if self.is_initialized:
            return
        
        try:
            logger.info(f"Loading embedding model: {self.model_name}")
            self.model = SentenceTransformer(self.model_name)
            self.dimension = self.model.get_sentence_embedding_dimension()
            self.is_initialized = True
            logger.info(f"Embedding model loaded, dimension: {self.dimension}")
        except Exception as e:
            logger.error(f"Failed to load embedding model: {e}")
            raise
    
    async def encode_text(self, text: str) -> List[float]:
        """Generate embedding for a single text"""
        if not self.is_initialized:
            await self.initialize()
        
        try:
            embedding = await asyncio.to_thread(self.model.encode, [text])
            return embedding[0].tolist()
        except Exception as e:
            logger.error(f"Embedding generation error: {e}")
            return [0.0] * self.dimension
    
    async def encode_texts(self, texts: List[str]) -> List[List[float]]:
        """Generate embeddings for multiple texts"""
        if not self.is_initialized:
            await self.initialize()
        
        try:
            embeddings = await asyncio.to_thread(self.model.encode, texts)
            return [emb.tolist() for emb in embeddings]
        except Exception as e:
            logger.error(f"Batch embedding generation error: {e}")
            return [[0.0] * self.dimension for _ in texts]

class DocumentProcessor:
    """Processes various document types for RAG"""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
            length_function=len,
        )
    
    async def process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Process PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_content = []
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    text_content.append(page_text)
                
                full_text = "\n\n".join(text_content)
                
                # Extract metadata
                metadata = {
                    'num_pages': len(pdf_reader.pages),
                    'title': pdf_reader.metadata.get('/Title', '') if pdf_reader.metadata else '',
                    'author': pdf_reader.metadata.get('/Author', '') if pdf_reader.metadata else '',
                    'creation_date': pdf_reader.metadata.get('/CreationDate', '') if pdf_reader.metadata else ''
                }
                
                return {
                    'content': full_text,
                    'metadata': metadata,
                    'page_contents': text_content
                }
        except Exception as e:
            logger.error(f"PDF processing error: {e}")
            raise
    
    async def process_docx(self, file_path: str) -> Dict[str, Any]:
        """Process DOCX file"""
        try:
            doc = docx.Document(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            full_text = "\n".join(text_content)
            
            # Basic metadata
            metadata = {
                'num_paragraphs': len([p for p in doc.paragraphs if p.text.strip()]),
                'title': doc.core_properties.title or '',
                'author': doc.core_properties.author or '',
                'created': str(doc.core_properties.created) if doc.core_properties.created else ''
            }
            
            return {
                'content': full_text,
                'metadata': metadata
            }
        except Exception as e:
            logger.error(f"DOCX processing error: {e}")
            raise
    
    async def process_text_file(self, file_path: str) -> Dict[str, Any]:
        """Process plain text file"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                content = await file.read()
            
            return {
                'content': content,
                'metadata': {
                    'file_size': len(content),
                    'line_count': len(content.splitlines())
                }
            }
        except Exception as e:
            logger.error(f"Text file processing error: {e}")
            raise
    
    def create_chunks(self, text: str) -> List[str]:
        """Split text into chunks"""
        try:
            documents = [Document(page_content=text)]
            chunks = self.text_splitter.split_documents(documents)
            return [chunk.page_content for chunk in chunks]
        except Exception as e:
            logger.error(f"Text chunking error: {e}")
            return [text]  # Return original text if chunking fails

class WebScraper:
    """Scrapes web content for RAG"""
    
    def __init__(self):
        self.session = None
    
    async def _get_session(self):
        """Get or create aiohttp session"""
        if not self.session:
            self.session = aiohttp.ClientSession(
                timeout=aiohttp.ClientTimeout(total=30),
                headers={'User-Agent': 'Mozilla/5.0 (compatible; ChatBot/1.0)'}
            )
        return self.session
    
    async def scrape_url(self, url: str) -> Dict[str, Any]:
        """Scrape content from a URL"""
        session = await self._get_session()
        
        try:
            async with session.get(url) as response:
                if response.status != 200:
                    raise Exception(f"HTTP {response.status}")
                
                content = await response.text()
                content_type = response.headers.get('content-type', '')
                
                # Parse HTML content
                soup = BeautifulSoup(content, 'html.parser')
                
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                
                # Extract text content
                text_content = soup.get_text()
                
                # Clean up text
                lines = (line.strip() for line in text_content.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = ' '.join(chunk for chunk in chunks if chunk)
                
                # Extract metadata
                title = soup.find('title').get_text() if soup.find('title') else ''
                meta_description = ''
                
                meta_desc = soup.find('meta', attrs={'name': 'description'})
                if meta_desc:
                    meta_description = meta_desc.get('content', '')
                
                # Extract keywords
                keywords = []
                meta_keywords = soup.find('meta', attrs={'name': 'keywords'})
                if meta_keywords:
                    keywords = [kw.strip() for kw in meta_keywords.get('content', '').split(',')]
                
                return {
                    'url': url,
                    'title': title,
                    'content': text,
                    'meta_description': meta_description,
                    'keywords': keywords,
                    'content_type': content_type,
                    'status_code': response.status,
                    'scraped_at': datetime.utcnow()
                }
                
        except Exception as e:
            logger.error(f"Web scraping error for {url}: {e}")
            raise
    
    async def scrape_multiple_urls(self, urls: List[str]) -> List[Dict[str, Any]]:
        """Scrape multiple URLs"""
        results = []
        for url in urls:
            try:
                result = await self.scrape_url(url)
                results.append(result)
                await asyncio.sleep(settings.SCRAPING_DELAY)  # Rate limiting
            except Exception as e:
                logger.error(f"Failed to scrape {url}: {e}")
                results.append({'url': url, 'error': str(e)})
        
        return results
    
    async def close(self):
        """Close the session"""
        if self.session:
            await self.session.close()

class VectorStore:
    """Vector store for similarity search using FAISS"""
    
    def __init__(self, dimension: int):
        self.dimension = dimension
        self.index = None
        self.documents = []
        self.is_initialized = False
    
    def initialize(self):
        """Initialize FAISS index"""
        if not self.is_initialized:
            self.index = faiss.IndexFlatIP(self.dimension)  # Inner product (cosine similarity)
            self.is_initialized = True
    
    def add_vectors(self, embeddings: List[List[float]], documents: List[Dict[str, Any]]):
        """Add vectors to the index"""
        if not self.is_initialized:
            self.initialize()
        
        # Normalize vectors for cosine similarity
        embeddings_np = np.array(embeddings).astype('float32')
        faiss.normalize_L2(embeddings_np)
        
        self.index.add(embeddings_np)
        self.documents.extend(documents)
    
    def search(self, query_embedding: List[float], k: int = 5) -> List[Tuple[Dict[str, Any], float]]:
        """Search for similar vectors"""
        if not self.is_initialized or self.index.ntotal == 0:
            return []
        
        # Normalize query vector
        query_np = np.array([query_embedding]).astype('float32')
        faiss.normalize_L2(query_np)
        
        scores, indices = self.index.search(query_np, min(k, self.index.ntotal))
        
        results = []
        for i, (score, idx) in enumerate(zip(scores[0], indices[0])):
            if idx < len(self.documents):
                results.append((self.documents[idx], float(score)))
        
        return results

class RAGSystem:
    """Complete RAG system"""
    
    def __init__(self):
        self.embedding_manager = EmbeddingManager()
        self.document_processor = DocumentProcessor()
        self.web_scraper = WebScraper()
        self.vector_store = VectorStore(settings.VECTOR_DIMENSION)
        self.is_initialized = False
    
    async def initialize(self):
        """Initialize the RAG system"""
        if self.is_initialized:
            return
        
        await self.embedding_manager.initialize()
        self.vector_store.dimension = self.embedding_manager.dimension
        self.vector_store.initialize()
        
        # Load existing documents from database
        await self._load_existing_documents()
        
        self.is_initialized = True
        logger.info("RAG system initialized successfully")
    
    async def _load_existing_documents(self):
        """Load existing documents from database into vector store"""
        try:
            with get_db_session() as db:
                # Load document chunks
                doc_chunks = db.query(DocumentChunk).filter(
                    DocumentChunk.embedding.isnot(None)
                ).all()
                
                if doc_chunks:
                    embeddings = []
                    documents = []
                    
                    for chunk in doc_chunks:
                        if chunk.embedding:
                            embeddings.append(chunk.embedding)
                            documents.append({
                                'id': chunk.id,
                                'content': chunk.content,
                                'source': 'document',
                                'document_id': chunk.document_id,
                                'chunk_index': chunk.chunk_index
                            })
                    
                    if embeddings:
                        self.vector_store.add_vectors(embeddings, documents)
                        logger.info(f"Loaded {len(embeddings)} document chunks")
                
                # Load web chunks
                web_chunks = db.query(WebChunk).filter(
                    WebChunk.embedding.isnot(None)
                ).all()
                
                if web_chunks:
                    embeddings = []
                    documents = []
                    
                    for chunk in web_chunks:
                        if chunk.embedding:
                            embeddings.append(chunk.embedding)
                            documents.append({
                                'id': chunk.id,
                                'content': chunk.content,
                                'source': 'web',
                                'source_id': chunk.source_id,
                                'chunk_index': chunk.chunk_index
                            })
                    
                    if embeddings:
                        self.vector_store.add_vectors(embeddings, documents)
                        logger.info(f"Loaded {len(embeddings)} web chunks")
                
                # Load knowledge base entries
                kb_entries = db.query(KnowledgeBase).filter(
                    KnowledgeBase.embedding.isnot(None)
                ).all()
                
                if kb_entries:
                    embeddings = []
                    documents = []
                    
                    for entry in kb_entries:
                        if entry.embedding:
                            embeddings.append(entry.embedding)
                            documents.append({
                                'id': entry.id,
                                'content': f"Q: {entry.question}\nA: {entry.answer}",
                                'source': 'knowledge_base',
                                'question': entry.question,
                                'answer': entry.answer,
                                'category': entry.category
                            })
                    
                    if embeddings:
                        self.vector_store.add_vectors(embeddings, documents)
                        logger.info(f"Loaded {len(embeddings)} knowledge base entries")
                
        except Exception as e:
            logger.error(f"Error loading existing documents: {e}")
    
    async def add_document(self, file_path: str, filename: str, file_type: str) -> str:
        """Add a document to the RAG system"""
        if not self.is_initialized:
            await self.initialize()
        
        try:
            # Process document based on type
            if file_type == 'pdf':
                doc_data = await self.document_processor.process_pdf(file_path)
            elif file_type == 'docx':
                doc_data = await self.document_processor.process_docx(file_path)
            elif file_type in ['txt', 'md']:
                doc_data = await self.document_processor.process_text_file(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
            
            # Create chunks
            chunks = self.document_processor.create_chunks(doc_data['content'])
            
            # Generate embeddings
            embeddings = await self.embedding_manager.encode_texts(chunks)
            
            # Save to database
            with get_db_session() as db:
                # Create document record
                db_doc = DBDocument(
                    filename=filename,
                    original_name=filename,
                    file_type=file_type,
                    file_size=os.path.getsize(file_path),
                    file_path=file_path,
                    content=doc_data['content'],
                    title=doc_data['metadata'].get('title', ''),
                    author=doc_data['metadata'].get('author', ''),
                    extra_metadata=doc_data['metadata'],
                    processing_status='completed',
                    processed_at=datetime.utcnow()
                )
                db.add(db_doc)
                db.flush()
                
                # Create chunk records
                for i, (chunk_content, embedding) in enumerate(zip(chunks, embeddings)):
                    db_chunk = DocumentChunk(
                        document_id=db_doc.id,
                        content=chunk_content,
                        chunk_index=i,
                        embedding=embedding,
                        embedding_model=self.embedding_manager.model_name
                    )
                    db.add(db_chunk)
                
                db.commit()
                
                # Add to vector store
                documents = []
                for i, chunk_content in enumerate(chunks):
                    documents.append({
                        'id': f"{db_doc.id}_{i}",
                        'content': chunk_content,
                        'source': 'document',
                        'document_id': db_doc.id,
                        'chunk_index': i
                    })
                
                self.vector_store.add_vectors(embeddings, documents)
                
                logger.info(f"Added document {filename} with {len(chunks)} chunks")
                return db_doc.id
                
        except Exception as e:
            logger.error(f"Error adding document: {e}")
            raise
    
    async def add_web_content(self, urls: List[str]) -> List[str]:
        """Add web content to the RAG system"""
        if not self.is_initialized:
            await self.initialize()
        
        results = []
        scraped_data = await self.web_scraper.scrape_multiple_urls(urls)
        
        for data in scraped_data:
            if 'error' in data:
                continue
            
            try:
                # Create chunks
                chunks = self.document_processor.create_chunks(data['content'])
                
                # Generate embeddings
                embeddings = await self.embedding_manager.encode_texts(chunks)
                
                # Save to database
                with get_db_session() as db:
                    # Check if URL already exists
                    existing = db.query(WebSource).filter(WebSource.url == data['url']).first()
                    if existing:
                        logger.info(f"URL {data['url']} already exists, skipping")
                        continue
                    
                    # Create web source record
                    parsed_url = urlparse(data['url'])
                    web_source = WebSource(
                        url=data['url'],
                        domain=parsed_url.netloc,
                        title=data['title'],
                        content=data['content'],
                        scraped_at=data['scraped_at'],
                        status_code=data['status_code'],
                        content_type=data['content_type'],
                        meta_description=data['meta_description'],
                        keywords=data['keywords']
                    )
                    db.add(web_source)
                    db.flush()
                    
                    # Create chunk records
                    for i, (chunk_content, embedding) in enumerate(zip(chunks, embeddings)):
                        web_chunk = WebChunk(
                            source_id=web_source.id,
                            content=chunk_content,
                            chunk_index=i,
                            embedding=embedding,
                            embedding_model=self.embedding_manager.model_name
                        )
                        db.add(web_chunk)
                    
                    db.commit()
                    
                    # Add to vector store
                    documents = []
                    for i, chunk_content in enumerate(chunks):
                        documents.append({
                            'id': f"{web_source.id}_{i}",
                            'content': chunk_content,
                            'source': 'web',
                            'source_id': web_source.id,
                            'chunk_index': i,
                            'url': data['url'],
                            'title': data['title']
                        })
                    
                    self.vector_store.add_vectors(embeddings, documents)
                    results.append(web_source.id)
                    
                    logger.info(f"Added web content from {data['url']} with {len(chunks)} chunks")
                    
            except Exception as e:
                logger.error(f"Error adding web content from {data['url']}: {e}")
        
        return results
    
    async def query(self, query: str, k: int = None, threshold: float = None) -> Dict[str, Any]:
        """Query the RAG system"""
        if not self.is_initialized:
            await self.initialize()
        
        k = k or settings.MAX_RETRIEVAL_DOCS
        threshold = threshold or settings.SIMILARITY_THRESHOLD
        
        try:
            # Generate query embedding
            query_embedding = await self.embedding_manager.encode_text(query)
            
            # Search vector store
            results = self.vector_store.search(query_embedding, k=k*2)  # Get more results to filter
            
            # Filter by threshold
            filtered_results = [(doc, score) for doc, score in results if score >= threshold]
            filtered_results = filtered_results[:k]  # Take top k
            
            # Generate response using LLM if we have relevant documents
            response = ""
            if filtered_results:
                # Prepare context from retrieved documents
                context_parts = []
                for doc, score in filtered_results:
                    context_parts.append(f"Source: {doc.get('source', 'unknown')}")
                    context_parts.append(f"Content: {doc['content']}")
                    context_parts.append(f"Relevance: {score:.3f}")
                    context_parts.append("---")
                
                context = "\n".join(context_parts)
                
                # Generate response using LLM
                messages = [
                    {
                        "role": "system", 
                        "content": "You are a helpful assistant. Use the provided context to answer the user's question. If the context doesn't contain relevant information, say so."
                    },
                    {
                        "role": "user", 
                        "content": f"Context:\n{context}\n\nQuestion: {query}"
                    }
                ]
                
                response = await llm_manager.generate_response(messages)
            else:
                response = "I couldn't find relevant information in the knowledge base to answer your question."
            
            return {
                'query': query,
                'response': response,
                'retrieved_documents': [
                    {
                        'content': doc['content'][:500] + '...' if len(doc['content']) > 500 else doc['content'],
                        'source': doc['source'],
                        'score': score,
                        'metadata': {k: v for k, v in doc.items() if k not in ['content', 'source']}
                    }
                    for doc, score in filtered_results
                ],
                'total_results': len(results),
                'filtered_results': len(filtered_results)
            }
            
        except Exception as e:
            logger.error(f"RAG query error: {e}")
            return {
                'query': query,
                'response': f"An error occurred while processing your query: {str(e)}",
                'retrieved_documents': [],
                'total_results': 0,
                'filtered_results': 0,
                'error': str(e)
            }
    
    async def add_knowledge_entry(self, question: str, answer: str, category: str = None, tags: List[str] = None) -> str:
        """Add a knowledge base entry"""
        if not self.is_initialized:
            await self.initialize()
        
        try:
            # Generate embedding for the question
            content = f"Q: {question}\nA: {answer}"
            embedding = await self.embedding_manager.encode_text(content)
            
            # Save to database
            with get_db_session() as db:
                kb_entry = KnowledgeBase(
                    question=question,
                    answer=answer,
                    category=category,
                    tags=tags or [],
                    embedding=embedding,
                    embedding_model=self.embedding_manager.model_name
                )
                db.add(kb_entry)
                db.commit()
                
                # Add to vector store
                document = {
                    'id': kb_entry.id,
                    'content': content,
                    'source': 'knowledge_base',
                    'question': question,
                    'answer': answer,
                    'category': category
                }
                
                self.vector_store.add_vectors([embedding], [document])
                
                logger.info(f"Added knowledge base entry: {question}")
                return kb_entry.id
                
        except Exception as e:
            logger.error(f"Error adding knowledge entry: {e}")
            raise
    
    async def get_statistics(self) -> Dict[str, Any]:
        """Get RAG system statistics"""
        try:
            with get_db_session() as db:
                doc_count = db.query(DBDocument).count()
                chunk_count = db.query(DocumentChunk).count()
                web_count = db.query(WebSource).count()
                web_chunk_count = db.query(WebChunk).count()
                kb_count = db.query(KnowledgeBase).count()
                
                return {
                    'documents': doc_count,
                    'document_chunks': chunk_count,
                    'web_sources': web_count,
                    'web_chunks': web_chunk_count,
                    'knowledge_base_entries': kb_count,
                    'total_vectors': self.vector_store.index.ntotal if self.vector_store.index else 0,
                    'embedding_model': self.embedding_manager.model_name,
                    'vector_dimension': self.embedding_manager.dimension
                }
        except Exception as e:
            logger.error(f"Error getting statistics: {e}")
            return {'error': str(e)}

# Global RAG system instance
rag_system = RAGSystem()
